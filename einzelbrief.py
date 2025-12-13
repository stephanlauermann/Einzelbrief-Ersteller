from __future__ import annotations

import datetime as dt
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


# -----------------------------
# DOCX helpers
# -----------------------------
def _mm_to_twips(mm: float) -> int:
    return int(mm / 25.4 * 1440)

def _mm_to_pt(mm: float) -> float:
    return mm * 72.0 / 25.4

def _mm_to_cm(mm: float) -> float:
    return mm / 10.0

def set_doc_default_arial(doc: Document, font_size_pt: int = 11):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

def add_paragraph_arial(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    align=None,
    size_pt: int | None = None,
    left_indent_mm: float | None = None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if left_indent_mm is not None:
        p.paragraph_format.left_indent = Cm(_mm_to_cm(left_indent_mm))
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.bold = bold
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    return p

def _set_row_height_mm(row, mm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(_mm_to_twips(mm)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)

def _set_table_left_indent_mm(table, mm: float):
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(_mm_to_twips(mm)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

def build_sender_line(first, last, street, plz, city) -> str:
    name = f"{first.strip()} {last.strip()}".strip()
    return f"{name} – {street.strip()} – {plz.strip()} {city.strip()}".strip()

def build_recipient_lines(first, last, street, extra, plz, city) -> list[str]:
    lines = []
    name = f"{first.strip()} {last.strip()}".strip()
    if name:
        lines.append(name)
    if street.strip():
        lines.append(street.strip())
    if extra.strip():
        lines.append(extra.strip())
    plz_city = f"{plz.strip()} {city.strip()}".strip()
    if plz_city:
        lines.append(plz_city)
    return lines

def dedupe_greeting_from_body(greeting: str, body: str) -> str:
    g = (greeting or "").strip()
    b = (body or "").lstrip()
    if g and b.startswith(g):
        lines = (body or "").splitlines()
        out, removed = [], False
        for ln in lines:
            if not removed and ln.strip() == g:
                removed = True
                continue
            out.append(ln)
        return "\n".join(out).lstrip("\n")
    return body


def add_window_area(doc: Document, sender_line: str, recipient_lines: list[str]) -> float:
    """
    Fensterbereich (links 20mm, oben 45mm). Gibt die linke Kante (mm) zurück,
    damit wir den Rest bündig einrücken können.
    """
    # Stellschrauben
    LEFT_FROM_PAGE_MM = 20.0
    TOP_TO_FIELD_MM = 45.0
    SENDER_ZONE_MM = 10.0
    RECIPIENT_ZONE_MM = 35.0
    AFTER_FIELD_MM = 10.0

    t = doc.add_table(rows=4, cols=1)
    t.autofit = False
    _remove_table_borders(t)
    _set_table_left_indent_mm(t, LEFT_FROM_PAGE_MM)

    t.columns[0].width = Cm(9.0)  # ~90mm

    _set_row_height_mm(t.rows[0], TOP_TO_FIELD_MM)
    _set_row_height_mm(t.rows[1], SENDER_ZONE_MM)
    _set_row_height_mm(t.rows[2], RECIPIENT_ZONE_MM)
    _set_row_height_mm(t.rows[3], AFTER_FIELD_MM)

    t.cell(0, 0).text = ""

    c1 = t.cell(1, 0)
    c1.text = ""
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(sender_line)
    r1.font.name = "Arial"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r1.font.size = Pt(8)

    c2 = t.cell(2, 0)
    c2.text = ""
    p2 = c2.paragraphs[0]
    for i, line in enumerate(recipient_lines):
        rr = p2.add_run(line)
        rr.font.name = "Arial"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if i < len(recipient_lines) - 1:
            p2.add_run("\n")

    t.cell(3, 0).text = ""
    return LEFT_FROM_PAGE_MM


def add_fold_and_punch_marks(doc: Document):
    X_LEFT_MM = 3.0
    LINE_LEN_MM = 8.0
    LINE_THICK_PT = 0.75
    FOLD_Y_MM = 99.0
    PUNCH_Y_MM = 148.5

    header = doc.sections[0].header

    def add_line(y_mm: float):
        x1_pt = _mm_to_pt(X_LEFT_MM)
        y_pt = _mm_to_pt(y_mm)
        x2_pt = _mm_to_pt(X_LEFT_MM + LINE_LEN_MM)

        line_xml = f"""
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:v="urn:schemas-microsoft-com:vml"
     xmlns:o="urn:schemas-microsoft-com:office:office">
  <w:r>
    <w:pict>
      <v:line from="{x1_pt}pt,{y_pt}pt"
              to="{x2_pt}pt,{y_pt}pt"
              style="position:absolute;left:{x1_pt}pt;top:{y_pt}pt;z-index:251659264"
              strokecolor="#000000"
              strokeweight="{LINE_THICK_PT}pt"/>
    </w:pict>
  </w:r>
</w:p>
"""
        header._element.append(parse_xml(line_xml))

    add_line(FOLD_Y_MM)
    add_line(PUNCH_Y_MM)


def make_letter_docx(data: dict, out_path: str):
    doc = Document()
    set_doc_default_arial(doc, font_size_pt=11)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    sender_line = build_sender_line(
        data["sender_first"], data["sender_last"],
        data["sender_street"], data["sender_plz"], data["sender_city"]
    )
    recipient_lines = build_recipient_lines(
        data["recipient_first"], data["recipient_last"],
        data["recipient_street"], data["recipient_extra"],
        data["recipient_plz"], data["recipient_city"]
    )

    body_left_mm = add_window_area(doc, sender_line, recipient_lines)

    # Datum (rechts, NICHT einrücken)
    date_line = f'{data["letter_city"].strip()}, den {data["letter_date"].strftime("%d.%m.%Y")}'.strip()
    add_paragraph_arial(doc, date_line, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_arial(doc, "")

    # Ab hier einrücken auf Fensterkante
    add_paragraph_arial(doc, data["subject"].strip(), bold=True, left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)

    greeting = data["greeting"].strip() or "Hallo,"
    add_paragraph_arial(doc, greeting, left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)

    body = dedupe_greeting_from_body(greeting, (data["body"] or "").rstrip("\n"))
    for line in body.splitlines():
        add_paragraph_arial(doc, line, left_indent_mm=body_left_mm)

    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, data["closing"].strip() or "Mit freundlichen Grüßen", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, data["signature_name"].strip(), left_indent_mm=body_left_mm)

    add_fold_and_punch_marks(doc)
    doc.save(out_path)


# -----------------------------
# GUI
# -----------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Einzelbrief-Ersteller")
        self.geometry("980x720")

        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)

        self.sender = self._sender_fields(self, row=0, col=0)
        self.recipient = self._recipient_fields(self, row=0, col=1)
        self._build_letter_fields()

        ttk.Button(self, text="DOCX exportieren…", command=self.export_docx).grid(
            row=3, column=0, columnspan=2, pady=12
        )

    def _sender_fields(self, parent, row, col):
        lf = ttk.LabelFrame(parent, text="Absender (Fenster oben klein)")
        lf.grid(row=row, column=col, sticky="nsew", padx=10, pady=10)

        fields = {}
        for i, (lab, key) in enumerate([
            ("Vorname", "sender_first"),
            ("Nachname", "sender_last"),
            ("Straße + Nr", "sender_street"),
            ("PLZ", "sender_plz"),
            ("Ort", "sender_city"),
        ]):
            ttk.Label(lf, text=lab).grid(row=i, column=0, sticky="w", padx=6, pady=4)
            e = ttk.Entry(lf, width=42)
            e.grid(row=i, column=1, sticky="w", padx=6, pady=4)
            fields[key] = e
        return fields

    def _recipient_fields(self, parent, row, col):
        lf = ttk.LabelFrame(parent, text="Empfänger (Fenster unten)")
        lf.grid(row=row, column=col, sticky="nsew", padx=10, pady=10)

        fields = {}
        for i, (lab, key) in enumerate([
            ("Vorname", "recipient_first"),
            ("Nachname", "recipient_last"),
            ("Straße + Nr", "recipient_street"),
            ("Zusatz (optional)", "recipient_extra"),
            ("PLZ", "recipient_plz"),
            ("Ort", "recipient_city"),
        ]):
            ttk.Label(lf, text=lab).grid(row=i, column=0, sticky="w", padx=6, pady=4)
            e = ttk.Entry(lf, width=42)
            e.grid(row=i, column=1, sticky="w", padx=6, pady=4)
            fields[key] = e
        return fields

    def _build_letter_fields(self):
        lf = ttk.LabelFrame(self, text="Briefdaten")
        lf.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)
        lf.columnconfigure(1, weight=1)
        lf.rowconfigure(3, weight=1)

        ttk.Label(lf, text="Ort (Datumszeile)").grid(row=0, column=0, sticky="w", padx=6, pady=4)
        self.letter_city = ttk.Entry(lf, width=30)
        self.letter_city.grid(row=0, column=1, sticky="w", padx=6, pady=4)

        ttk.Label(lf, text="Datum (TT.MM.JJJJ)").grid(row=0, column=2, sticky="w", padx=6, pady=4)
        self.letter_date = ttk.Entry(lf, width=14)
        self.letter_date.grid(row=0, column=3, sticky="w", padx=6, pady=4)
        self.letter_date.insert(0, dt.date.today().strftime("%d.%m.%Y"))

        ttk.Label(lf, text="Betreff").grid(row=1, column=0, sticky="w", padx=6, pady=4)
        self.subject = ttk.Entry(lf)
        self.subject.grid(row=1, column=1, columnspan=3, sticky="ew", padx=6, pady=4)

        ttk.Label(lf, text="Anrede").grid(row=2, column=0, sticky="w", padx=6, pady=4)
        self.greeting = ttk.Entry(lf)
        self.greeting.grid(row=2, column=1, sticky="ew", padx=6, pady=4)
        self.greeting.insert(0, "Hallo,")

        ttk.Label(lf, text="Grußformel").grid(row=2, column=2, sticky="w", padx=6, pady=4)
        self.closing = ttk.Entry(lf)
        self.closing.grid(row=2, column=3, sticky="w", padx=6, pady=4)
        self.closing.insert(0, "Mit freundlichen Grüßen")

        ttk.Label(lf, text="Text (ohne Anrede)").grid(row=3, column=0, sticky="nw", padx=6, pady=4)
        self.body = tk.Text(lf, height=14)
        self.body.grid(row=3, column=1, columnspan=3, sticky="nsew", padx=6, pady=4)

        ttk.Label(lf, text="Signatur-Name").grid(row=4, column=0, sticky="w", padx=6, pady=4)
        self.signature = ttk.Entry(lf)
        self.signature.grid(row=4, column=1, sticky="w", padx=6, pady=4)

    def _collect_data(self) -> dict:
        d = {}
        for k, e in self.sender.items():
            d[k] = e.get().strip()
        for k, e in self.recipient.items():
            d[k] = e.get().strip()

        d["letter_city"] = self.letter_city.get().strip()
        try:
            d["letter_date"] = dt.datetime.strptime(self.letter_date.get().strip(), "%d.%m.%Y").date()
        except ValueError:
            raise ValueError("Datum bitte als TT.MM.JJJJ (z.B. 12.12.2025) eingeben.")

        d["subject"] = self.subject.get().strip()
        d["greeting"] = self.greeting.get().strip() or "Hallo,"
        d["body"] = self.body.get("1.0", "end").rstrip("\n")
        d["closing"] = self.closing.get().strip() or "Mit freundlichen Grüßen"

        sig = self.signature.get().strip()
        if not sig:
            sig = f"{d['sender_first']} {d['sender_last']}".strip()
        d["signature_name"] = sig

        required = [
            "sender_first", "sender_last", "sender_street", "sender_plz", "sender_city",
            "recipient_first", "recipient_last", "recipient_street", "recipient_plz", "recipient_city",
            "letter_city", "subject", "signature_name"
        ]
        for k in required:
            if not d.get(k):
                raise ValueError(f"Bitte Feld ausfüllen: {k}")

        return d

    def export_docx(self):
        try:
            data = self._collect_data()
        except Exception as e:
            messagebox.showerror("Fehler", str(e))
            return

        out = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Dokument", "*.docx")])
        if not out:
            return

        try:
            make_letter_docx(data, out)
            messagebox.showinfo("OK", f"Brief erstellt:\n{out}")
        except Exception as e:
            messagebox.showerror("Fehler", str(e))


if __name__ == "__main__":
    App().mainloop()
