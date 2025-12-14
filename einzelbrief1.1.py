from __future__ import annotations

import datetime as dt
import json
import threading
import time
import urllib.parse
import urllib.request
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import subprocess
import tempfile
import shutil
import os

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm as RL_MM
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ============================================================
# Einstellungen (Fenster / Marken / Autocomplete)
# ============================================================

# Fensterbrief-Startwerte (typisch “Amtsumschlag” DL/C6-5)
LEFT_FROM_PAGE_MM = 20.0
TOP_TO_FIELD_MM = 45.0
SENDER_ZONE_MM = 10.0
RECIPIENT_ZONE_MM = 35.0
AFTER_FIELD_MM = 10.0

# Striche links
FOLD_Y_MM = 99.0     # Falzmarke (Drittel-Falz)
PUNCH_Y_MM = 148.5   # Lochmarke (A4 Mitte)
MARK_X_LEFT_MM = 3.0
MARK_LEN_MM = 8.0
MARK_THICK_PT = 0.75

# OSM / Nominatim
NOMINATIM_URL = "https://nominatim.openstreetmap.org/search"
NOMINATIM_LIMIT = 8
AUTOCOMPLETE_DEBOUNCE_MS = 350
AUTOCOMPLETE_CACHE_MAX = 200

# Wichtig: Nominatim möchte einen sinnvollen User-Agent. Trage hier deinen Namen/Projekt ein.
NOMINATIM_USER_AGENT = "BriefErsteller/1.0 (contact: example@example.com)"


# ============================================================
# Hilfsfunktionen (Maße / Arial)
# ============================================================
def _mm_to_twips(mm: float) -> int:
    return int(mm / 25.4 * 1440)

def _mm_to_pt(mm: float) -> float:
    return mm * 72.0 / 25.4

def _mm_to_cm(mm: float) -> float:
    return mm / 10.0


# ============================================================
# DOCX-Export
# ============================================================
def set_doc_default_arial(doc: Document, font_size_pt: int = 11):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(font_size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

def add_paragraph_arial(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    align=None,
    size_pt: int | None = None,
    left_indent_mm: float | None = None,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if left_indent_mm is not None:
        p.paragraph_format.left_indent = Cm(_mm_to_cm(left_indent_mm))
    r = p.add_run(text)
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r.bold = bold
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    return p

def _set_row_height_mm(row, mm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(_mm_to_twips(mm)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)

def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)

def _set_table_left_indent_mm(table, mm: float):
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(_mm_to_twips(mm)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

def build_sender_line(first, last, street, plz, city) -> str:
    name = f"{first.strip()} {last.strip()}".strip()
    return f"{name} – {street.strip()} – {plz.strip()} {city.strip()}".strip()

def build_recipient_lines(first, last, street, extra, plz, city) -> list[str]:
    lines = []
    name = f"{first.strip()} {last.strip()}".strip()
    if name:
        lines.append(name)
    if street.strip():
        lines.append(street.strip())
    if extra.strip():
        lines.append(extra.strip())
    plz_city = f"{plz.strip()} {city.strip()}".strip()
    if plz_city:
        lines.append(plz_city)
    return lines

def dedupe_greeting_from_body(greeting: str, body: str) -> str:
    g = (greeting or "").strip()
    b = (body or "").lstrip()
    if g and b.startswith(g):
        lines = (body or "").splitlines()
        out, removed = [], False
        for ln in lines:
            if not removed and ln.strip() == g:
                removed = True
                continue
            out.append(ln)
        return "\n".join(out).lstrip("\n")
    return body

def add_window_area_docx(doc: Document, sender_line: str, recipient_lines: list[str]) -> float:
    t = doc.add_table(rows=4, cols=1)
    t.autofit = False
    _remove_table_borders(t)
    _set_table_left_indent_mm(t, LEFT_FROM_PAGE_MM)

    t.columns[0].width = Cm(9.0)  # ~90mm

    _set_row_height_mm(t.rows[0], TOP_TO_FIELD_MM)
    _set_row_height_mm(t.rows[1], SENDER_ZONE_MM)
    _set_row_height_mm(t.rows[2], RECIPIENT_ZONE_MM)
    _set_row_height_mm(t.rows[3], AFTER_FIELD_MM)

    t.cell(0, 0).text = ""

    c1 = t.cell(1, 0)
    c1.text = ""
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(sender_line)
    r1.font.name = "Arial"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r1.font.size = Pt(8)

    c2 = t.cell(2, 0)
    c2.text = ""
    p2 = c2.paragraphs[0]
    for i, line in enumerate(recipient_lines):
        rr = p2.add_run(line)
        rr.font.name = "Arial"
        rr._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        if i < len(recipient_lines) - 1:
            p2.add_run("\n")

    t.cell(3, 0).text = ""
    return LEFT_FROM_PAGE_MM

def add_fold_and_punch_marks_docx(doc: Document):
    header = doc.sections[0].header

    def add_line(y_mm: float):
        x1_pt = _mm_to_pt(MARK_X_LEFT_MM)
        y_pt = _mm_to_pt(y_mm)
        x2_pt = _mm_to_pt(MARK_X_LEFT_MM + MARK_LEN_MM)

        line_xml = f"""
<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:v="urn:schemas-microsoft-com:vml"
     xmlns:o="urn:schemas-microsoft-com:office:office">
  <w:r>
    <w:pict>
      <v:line from="{x1_pt}pt,{y_pt}pt"
              to="{x2_pt}pt,{y_pt}pt"
              style="position:absolute;left:{x1_pt}pt;top:{y_pt}pt;z-index:251659264"
              strokecolor="#000000"
              strokeweight="{MARK_THICK_PT}pt"/>
    </w:pict>
  </w:r>
</w:p>
"""
        header._element.append(parse_xml(line_xml))

    add_line(FOLD_Y_MM)
    add_line(PUNCH_Y_MM)

def export_docx(data: dict, out_path: str):
    doc = Document()
    set_doc_default_arial(doc, font_size_pt=11)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    sender_line = build_sender_line(
        data["sender_first"], data["sender_last"],
        data["sender_street"], data["sender_plz"], data["sender_city"]
    )
    recipient_lines = build_recipient_lines(
        data["recipient_first"], data["recipient_last"],
        data["recipient_street"], data["recipient_extra"],
        data["recipient_plz"], data["recipient_city"]
    )

    body_left_mm = add_window_area_docx(doc, sender_line, recipient_lines)

    date_line = f'{data["letter_city"].strip()}, den {data["letter_date"].strftime("%d.%m.%Y")}'.strip()
    add_paragraph_arial(doc, date_line, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph_arial(doc, "")

    add_paragraph_arial(doc, data["subject"].strip(), bold=True, left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)

    greeting = data["greeting"].strip() or "Hallo,"
    add_paragraph_arial(doc, greeting, left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)

    body = dedupe_greeting_from_body(greeting, (data["body"] or "").rstrip("\n"))
    for line in body.splitlines():
        add_paragraph_arial(doc, line, left_indent_mm=body_left_mm)

    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, data["closing"].strip() or "Mit freundlichen Grüßen", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, "", left_indent_mm=body_left_mm)
    add_paragraph_arial(doc, data["signature_name"].strip(), left_indent_mm=body_left_mm)

    add_fold_and_punch_marks_docx(doc)
    doc.save(out_path)


# ============================================================
# PDF-Export (ReportLab)
# ============================================================
def _try_register_arial():
    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        r"/Library/Fonts/Arial.ttf",
        r"/System/Library/Fonts/Supplemental/Arial.ttf",
        r"/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
        r"/usr/share/fonts/truetype/msttcorefonts/arial.ttf",
    ]
    for path in candidates:
        try:
            pdfmetrics.registerFont(TTFont("Arial", path))
            return "Arial"
        except Exception:
            continue
    return "Helvetica"

def export_pdf(data: dict, out_path: str):
    font_name = _try_register_arial()
    w, h = A4
    c = canvas.Canvas(out_path, pagesize=A4)
    c.setTitle("Brief")

    def x(mm_val: float) -> float:
        return mm_val * RL_MM

    def y_from_top(mm_from_top: float) -> float:
        return h - mm_from_top * RL_MM

    # Marken
    c.setLineWidth(MARK_THICK_PT)
    c.line(x(MARK_X_LEFT_MM), y_from_top(FOLD_Y_MM), x(MARK_X_LEFT_MM + MARK_LEN_MM), y_from_top(FOLD_Y_MM))
    c.line(x(MARK_X_LEFT_MM), y_from_top(PUNCH_Y_MM), x(MARK_X_LEFT_MM + MARK_LEN_MM), y_from_top(PUNCH_Y_MM))

    sender_line = build_sender_line(
        data["sender_first"], data["sender_last"],
        data["sender_street"], data["sender_plz"], data["sender_city"]
    )
    recipient_lines = build_recipient_lines(
        data["recipient_first"], data["recipient_last"],
        data["recipient_street"], data["recipient_extra"],
        data["recipient_plz"], data["recipient_city"]
    )

    # Absender klein
    c.setFont(font_name, 8)
    sender_baseline_mm = TOP_TO_FIELD_MM + 7.0
    c.drawString(x(LEFT_FROM_PAGE_MM), y_from_top(sender_baseline_mm), sender_line)

    # Empfänger
    c.setFont(font_name, 11)
    recipient_start_mm = TOP_TO_FIELD_MM + SENDER_ZONE_MM + 7.0
    line_step_mm = 5.0
    for i, line in enumerate(recipient_lines):
        c.drawString(x(LEFT_FROM_PAGE_MM), y_from_top(recipient_start_mm + i * line_step_mm), line)

    # Datum rechts
    date_line = f'{data["letter_city"].strip()}, den {data["letter_date"].strftime("%d.%m.%Y")}'.strip()
    date_y_mm = TOP_TO_FIELD_MM + SENDER_ZONE_MM + RECIPIENT_ZONE_MM + AFTER_FIELD_MM + 8.0
    c.setFont(font_name, 11)
    c.drawRightString(w - x(20.0), y_from_top(date_y_mm), date_line)

    # Briefteil eingerückt auf Fensterkante
    cur_mm = date_y_mm + 12.0
    left_mm = LEFT_FROM_PAGE_MM

    # Betreff (fett wenn möglich)
    if font_name == "Helvetica":
        c.setFont("Helvetica-Bold", 11)
    else:
        c.setFont(font_name, 11)
    c.drawString(x(left_mm), y_from_top(cur_mm), data["subject"].strip())
    cur_mm += 10.0

    # Anrede
    c.setFont(font_name, 11)
    greeting = data["greeting"].strip() or "Hallo,"
    c.drawString(x(left_mm), y_from_top(cur_mm), greeting)
    cur_mm += 10.0

    # Body
    body = dedupe_greeting_from_body(greeting, (data["body"] or "").rstrip("\n"))
    for line in body.splitlines():
        if not line.strip():
            cur_mm += 5.0
            continue
        c.drawString(x(left_mm), y_from_top(cur_mm), line)
        cur_mm += 5.0

    cur_mm += 8.0
    closing = data["closing"].strip() or "Mit freundlichen Grüßen"
    c.drawString(x(left_mm), y_from_top(cur_mm), closing)
    cur_mm += 12.0
    c.drawString(x(left_mm), y_from_top(cur_mm), data["signature_name"].strip())

    c.showPage()
    c.save()


# ============================================================
# ODT-Export (LibreOffice / soffice)
# ============================================================
def _find_soffice() -> str | None:
    # 1) PATH
    p = shutil.which("soffice")
    if p:
        return p

    # 2) Windows typische Installpfade
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if os.path.isfile(c):
            return c

    # 3) macOS
    mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.isfile(mac):
        return mac

    return None

def convert_docx_to_odt(docx_path: str, out_odt_path: str):
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) nicht gefunden.\n"
            "Bitte LibreOffice installieren, dann erneut versuchen."
        )

    out_dir = os.path.dirname(os.path.abspath(out_odt_path))
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    expected = os.path.join(out_dir, base_name + ".odt")

    # LibreOffice schreibt in outdir mit eigenem Namen -> wir konvertieren in out_dir und benennen ggf. um
    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nodefault",
        "--norestore",
        "--convert-to", "odt",
        "--outdir", out_dir,
        docx_path,
    ]

    proc = subprocess.run(cmd, capture_output=True, text=True)
    if proc.returncode != 0:
        msg = (proc.stderr or proc.stdout or "").strip()
        raise RuntimeError("ODT-Konvertierung fehlgeschlagen.\n\n" + (msg or "(keine Details)"))

    if not os.path.isfile(expected):
        # Manche LibreOffice-Versionen nehmen leicht andere Namen – versuchen wir zu finden
        for fn in os.listdir(out_dir):
            if fn.lower().endswith(".odt") and fn.lower().startswith(base_name.lower()):
                expected = os.path.join(out_dir, fn)
                break

    if not os.path.isfile(expected):
        raise RuntimeError("LibreOffice meldet Erfolg, aber keine ODT-Datei wurde gefunden.")

    # Wenn Zielname abweicht, umbenennen
    if os.path.abspath(expected) != os.path.abspath(out_odt_path):
        if os.path.exists(out_odt_path):
            os.remove(out_odt_path)
        os.replace(expected, out_odt_path)


# ============================================================
# OSM / Nominatim Autocomplete (Debounce + Cache)
# ============================================================
class LruCache:
    def __init__(self, max_size: int):
        self.max_size = max_size
        self._d: dict[str, tuple[float, object]] = {}

    def get(self, key: str):
        v = self._d.get(key)
        if not v:
            return None
        ts, val = v
        self._d.pop(key, None)
        self._d[key] = (time.time(), val)
        return val

    def set(self, key: str, val: object):
        if key in self._d:
            self._d.pop(key, None)
        self._d[key] = (time.time(), val)
        while len(self._d) > self.max_size:
            oldest = min(self._d.items(), key=lambda kv: kv[1][0])[0]
            self._d.pop(oldest, None)

OSM_CACHE = LruCache(AUTOCOMPLETE_CACHE_MAX)

def nominatim_search(query: str):
    query = query.strip()
    if not query:
        return []

    cached = OSM_CACHE.get(query.lower())
    if cached is not None:
        return cached

    params = {
        "format": "json",
        "q": query,
        "addressdetails": "1",
        "limit": str(NOMINATIM_LIMIT),
        "countrycodes": "de",  # wenn international: entfernen
    }
    url = NOMINATIM_URL + "?" + urllib.parse.urlencode(params)

    req = urllib.request.Request(
        url,
        headers={"User-Agent": NOMINATIM_USER_AGENT, "Accept": "application/json"},
        method="GET",
    )

    with urllib.request.urlopen(req, timeout=8) as resp:
        data = json.loads(resp.read().decode("utf-8"))

    results = []
    for item in data:
        addr = item.get("address", {}) or {}
        road = addr.get("road") or addr.get("pedestrian") or addr.get("footway") or ""
        house = addr.get("house_number") or ""
        postcode = addr.get("postcode") or ""
        city = addr.get("city") or addr.get("town") or addr.get("village") or addr.get("hamlet") or ""
        display = item.get("display_name") or ""

        pretty = ", ".join([p for p in [
            (" ".join([road, house]).strip() or "").strip(),
            (" ".join([postcode, city]).strip() or "").strip(),
        ] if p])

        results.append({
            "pretty": pretty or display,
            "road": road,
            "house_number": house,
            "postcode": postcode,
            "city": city,
            "display": display,
        })

    OSM_CACHE.set(query.lower(), results)
    return results


class SuggestPopup(tk.Toplevel):
    def __init__(self, master, on_pick):
        super().__init__(master)
        self.withdraw()
        self.overrideredirect(True)
        self.listbox = tk.Listbox(self, height=8)
        self.listbox.pack(fill="both", expand=True)
        self.on_pick = on_pick
        self._items = []

        self.listbox.bind("<ButtonRelease-1>", self._pick)
        self.listbox.bind("<Return>", self._pick)
        self.listbox.bind("<Escape>", lambda e: self.hide())

    def show(self, x, y, width, items):
        self._items = items
        self.listbox.delete(0, "end")
        for it in items:
            self.listbox.insert("end", it["pretty"])
        self.geometry(f"{max(width, 320)}x180+{x}+{y}")
        self.deiconify()
        self.lift()

    def hide(self):
        self.withdraw()

    def _pick(self, _event=None):
        sel = self.listbox.curselection()
        if not sel:
            return
        item = self._items[sel[0]]
        self.hide()
        self.on_pick(item)


# ============================================================
# GUI
# ============================================================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Einzelbrief-Ersteller (OSM + DOCX/PDF/ODT)")
        self.geometry("1040x780")

        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=1)
        self.rowconfigure(2, weight=1)

        self._debounce_job = None
        self._active_entry = None

        self.suggest = SuggestPopup(self, self._apply_suggestion)

        self.sender = self._sender_fields(self, row=0, col=0)
        self.recipient = self._recipient_fields(self, row=0, col=1)
        self._build_letter_fields()

        btn_row = ttk.Frame(self)
        btn_row.grid(row=3, column=0, columnspan=2, pady=12)

        ttk.Button(btn_row, text="DOCX exportieren…", command=self.export_docx_btn).pack(side="left", padx=8)
        ttk.Button(btn_row, text="PDF exportieren…", command=self.export_pdf_btn).pack(side="left", padx=8)
        ttk.Button(btn_row, text="ODT exportieren…", command=self.export_odt_btn).pack(side="left", padx=8)

        self.bind_all("<Button-1>", self._global_click, add="+")

    def _sender_fields(self, parent, row, col):
        lf = ttk.LabelFrame(parent, text="Absender (Fenster oben klein)")
        lf.grid(row=row, column=col, sticky="nsew", padx=10, pady=10)

        fields = {}
        def mk(label, key, r):
            ttk.Label(lf, text=label).grid(row=r, column=0, sticky="w", padx=6, pady=4)
            e = ttk.Entry(lf, width=46)
            e.grid(row=r, column=1, sticky="w", padx=6, pady=4)
            fields[key] = e
            return e

        mk("Vorname", "sender_first", 0)
        mk("Nachname", "sender_last", 1)
        e_street = mk("Straße + Nr", "sender_street", 2)
        e_plz = mk("PLZ", "sender_plz", 3)
        e_city = mk("Ort", "sender_city", 4)

        self._attach_autocomplete(e_street, plz_entry=e_plz, city_entry=e_city)
        return fields

    def _recipient_fields(self, parent, row, col):
        lf = ttk.LabelFrame(parent, text="Empfänger (Fenster unten)")
        lf.grid(row=row, column=col, sticky="nsew", padx=10, pady=10)

        fields = {}
        def mk(label, key, r):
            ttk.Label(lf, text=label).grid(row=r, column=0, sticky="w", padx=6, pady=4)
            e = ttk.Entry(lf, width=46)
            e.grid(row=r, column=1, sticky="w", padx=6, pady=4)
            fields[key] = e
            return e

        mk("Vorname", "recipient_first", 0)
        mk("Nachname", "recipient_last", 1)
        e_street = mk("Straße + Nr", "recipient_street", 2)
        mk("Zusatz (optional)", "recipient_extra", 3)
        e_plz = mk("PLZ", "recipient_plz", 4)
        e_city = mk("Ort", "recipient_city", 5)

        self._attach_autocomplete(e_street, plz_entry=e_plz, city_entry=e_city)
        return fields

    def _attach_autocomplete(self, entry: ttk.Entry, *, plz_entry: ttk.Entry, city_entry: ttk.Entry):
        def on_key(_evt=None):
            self._active_entry = (entry, plz_entry, city_entry)
            if self._debounce_job:
                self.after_cancel(self._debounce_job)
            self._debounce_job = self.after(AUTOCOMPLETE_DEBOUNCE_MS, self._do_autocomplete)

        entry.bind("<KeyRelease>", on_key)
        entry.bind("<Down>", lambda e: self._focus_popup())

    def _do_autocomplete(self):
        if not self._active_entry:
            return
        entry, _, _ = self._active_entry
        q = entry.get().strip()
        if len(q) < 4:
            self.suggest.hide()
            return

        def worker(query: str, token: int):
            try:
                res = nominatim_search(query)
            except Exception:
                res = []
            self.after(0, lambda: self._show_suggestions_if_current(token, res))

        token = int(time.time() * 1000)
        entry._autocomplete_token = token  # type: ignore[attr-defined]
        threading.Thread(target=worker, args=(q, token), daemon=True).start()

    def _show_suggestions_if_current(self, token: int, results: list[dict]):
        if not self._active_entry:
            return
        entry, *_ = self._active_entry
        if getattr(entry, "_autocomplete_token", None) != token:
            return
        if not results:
            self.suggest.hide()
            return
        x = entry.winfo_rootx()
        y = entry.winfo_rooty() + entry.winfo_height()
        w = entry.winfo_width()
        self.suggest.show(x, y, w, results)

    def _apply_suggestion(self, item: dict):
        if not self._active_entry:
            return
        entry, plz_entry, city_entry = self._active_entry

        road = (item.get("road") or "").strip()
        house = (item.get("house_number") or "").strip()
        street = " ".join([road, house]).strip() or item.get("pretty", "")

        entry.delete(0, "end")
        entry.insert(0, street)

        if item.get("postcode"):
            plz_entry.delete(0, "end")
            plz_entry.insert(0, item["postcode"])
        if item.get("city"):
            city_entry.delete(0, "end")
            city_entry.insert(0, item["city"])

    def _focus_popup(self):
        try:
            if self.suggest.state() == "normal":
                self.suggest.listbox.focus_set()
                self.suggest.listbox.selection_set(0)
        except Exception:
            pass

    def _global_click(self, event):
        w = event.widget
        if w is self.suggest.listbox:
            return
        if isinstance(w, ttk.Entry) or isinstance(w, tk.Entry):
            return
        self.suggest.hide()

    def _build_letter_fields(self):
        lf = ttk.LabelFrame(self, text="Briefdaten")
        lf.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)
        lf.columnconfigure(1, weight=1)
        lf.rowconfigure(3, weight=1)

        ttk.Label(lf, text="Ort (Datumszeile)").grid(row=0, column=0, sticky="w", padx=6, pady=4)
        self.letter_city = ttk.Entry(lf, width=30)
        self.letter_city.grid(row=0, column=1, sticky="w", padx=6, pady=4)

        ttk.Label(lf, text="Datum (TT.MM.JJJJ)").grid(row=0, column=2, sticky="w", padx=6, pady=4)
        self.letter_date = ttk.Entry(lf, width=14)
        self.letter_date.grid(row=0, column=3, sticky="w", padx=6, pady=4)
        self.letter_date.insert(0, dt.date.today().strftime("%d.%m.%Y"))

        ttk.Label(lf, text="Betreff").grid(row=1, column=0, sticky="w", padx=6, pady=4)
        self.subject = ttk.Entry(lf)
        self.subject.grid(row=1, column=1, columnspan=3, sticky="ew", padx=6, pady=4)

        ttk.Label(lf, text="Anrede").grid(row=2, column=0, sticky="w", padx=6, pady=4)
        self.greeting = ttk.Entry(lf)
        self.greeting.grid(row=2, column=1, sticky="ew", padx=6, pady=4)
        self.greeting.insert(0, "Hallo,")

        ttk.Label(lf, text="Grußformel").grid(row=2, column=2, sticky="w", padx=6, pady=4)
        self.closing = ttk.Entry(lf)
        self.closing.grid(row=2, column=3, sticky="w", padx=6, pady=4)
        self.closing.insert(0, "Mit freundlichen Grüßen")

        ttk.Label(lf, text="Text (ohne Anrede)").grid(row=3, column=0, sticky="nw", padx=6, pady=4)
        self.body = tk.Text(lf, height=14)
        self.body.grid(row=3, column=1, columnspan=3, sticky="nsew", padx=6, pady=4)

        ttk.Label(lf, text="Signatur-Name").grid(row=4, column=0, sticky="w", padx=6, pady=4)
        self.signature = ttk.Entry(lf)
        self.signature.grid(row=4, column=1, sticky="w", padx=6, pady=4)

    def _collect_data(self) -> dict:
        d = {}
        for k, e in self.sender.items():
            d[k] = e.get().strip()
        for k, e in self.recipient.items():
            d[k] = e.get().strip()

        d["letter_city"] = self.letter_city.get().strip()
        try:
            d["letter_date"] = dt.datetime.strptime(self.letter_date.get().strip(), "%d.%m.%Y").date()
        except ValueError:
            raise ValueError("Datum bitte als TT.MM.JJJJ (z.B. 12.12.2025) eingeben.")

        d["subject"] = self.subject.get().strip()
        d["greeting"] = self.greeting.get().strip() or "Hallo,"
        d["body"] = self.body.get("1.0", "end").rstrip("\n")
        d["closing"] = self.closing.get().strip() or "Mit freundlichen Grüßen"

        sig = self.signature.get().strip()
        if not sig:
            sig = f"{d['sender_first']} {d['sender_last']}".strip()
        d["signature_name"] = sig

        required = [
            "sender_first", "sender_last", "sender_street", "sender_plz", "sender_city",
            "recipient_first", "recipient_last", "recipient_street", "recipient_plz", "recipient_city",
            "letter_city", "subject", "signature_name"
        ]
        for k in required:
            if not d.get(k):
                raise ValueError(f"Bitte Feld ausfüllen: {k}")

        return d

    # -------------------------
    # Export buttons
    # -------------------------
    def export_docx_btn(self):
        try:
            data = self._collect_data()
        except Exception as e:
            messagebox.showerror("Fehler", str(e))
            return

        out = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Dokument", "*.docx")])
        if not out:
            return
        try:
            export_docx(data, out)
            messagebox.showinfo("OK", f"DOCX erstellt:\n{out}")
        except Exception as e:
            messagebox.showerror("Fehler", str(e))

    def export_pdf_btn(self):
        try:
            data = self._collect_data()
        except Exception as e:
            messagebox.showerror("Fehler", str(e))
            return

        out = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
        if not out:
            return
        try:
            export_pdf(data, out)
            messagebox.showinfo("OK", f"PDF erstellt:\n{out}")
        except Exception as e:
            messagebox.showerror("Fehler", str(e))

    def export_odt_btn(self):
        try:
            data = self._collect_data()
        except Exception as e:
            messagebox.showerror("Fehler", str(e))
            return

        out_odt = filedialog.asksaveasfilename(defaultextension=".odt", filetypes=[("OpenDocument Text", "*.odt")])
        if not out_odt:
            return

        # DOCX temporär erzeugen und dann konvertieren
        try:
            with tempfile.TemporaryDirectory() as td:
                tmp_docx = os.path.join(td, "brief.docx")
                export_docx(data, tmp_docx)
                convert_docx_to_odt(tmp_docx, out_odt)
            messagebox.showinfo("OK", f"ODT erstellt:\n{out_odt}")
        except Exception as e:
            messagebox.showerror("Fehler", str(e))


if __name__ == "__main__":
    App().mainloop()
